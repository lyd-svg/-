"""
文档解析器
支持 PDF / DOCX / XLSX / TXT / MD / CSV 格式，
自动检测 QA 问答对格式，文本分块。
"""
import os
import re
from .config import CHUNK_SIZE, CHUNK_OVERLAP

# ========== 文档解析器 ==========


def parse_pdf(file_path: str) -> tuple[str, list[str]]:
    """解析 PDF 文件，返回 (全文文本, 段落列表)"""
    import pdfplumber
    text_parts = []
    paragraphs = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
            for para in page_text.split("\n\n"):
                para = para.strip()
                if para:
                    paragraphs.append(para)
    return "\n\n".join(text_parts), paragraphs


def parse_docx(file_path: str) -> tuple[str, list[str]]:
    """解析 DOCX 文件（段落 + 表格）"""
    import docx
    doc = docx.Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append("| " + " | ".join(cells) + " |")
        if rows_data:
            col_count = max(len(row.cells) for row in table.rows)
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            table_text = "\n".join([rows_data[0], separator] + rows_data[1:])
            paragraphs.append(table_text)

    return "\n\n".join(paragraphs), paragraphs


def parse_xlsx(file_path: str) -> tuple[str, list[str]]:
    """解析 XLSX 文件，每个工作表转为 Markdown 表格"""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    all_text = []
    paragraphs = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = "| " + " | ".join(str(c) if c is not None else "" for c in rows[0]) + " |"
        sep = "| " + " | ".join(["---"] * len(rows[0])) + " |"
        data_rows = []
        for row in rows[1:]:
            data_rows.append("| " + " | ".join(str(c) if c is not None else "" for c in row) + " |")
        table = f"## 工作表: {sheet_name}\n\n{header}\n{sep}\n" + "\n".join(data_rows[:50])
        paragraphs.append(table)
        all_text.append(f"工作表 {sheet_name}:\n" + "\n".join(
            str(c) for row in rows for c in row if c is not None
        ))
    wb.close()
    return "\n\n".join(all_text), paragraphs


def parse_txt(file_path: str) -> tuple[str, list[str]]:
    """解析 TXT 文件"""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    return text, paragraphs


def parse_md(file_path: str) -> tuple[str, list[str]]:
    """解析 Markdown 文件（与 TXT 相同处理）"""
    return parse_txt(file_path)


def parse_csv(file_path: str) -> tuple[str, list[str]]:
    """解析 CSV 文件，转为 Markdown 表格"""
    import pandas as pd
    df = pd.read_csv(file_path)
    text = df.to_markdown(index=False, numalign="left")
    paragraphs = [text]
    return text, paragraphs


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".txt": parse_txt,
    ".md": parse_md,
    ".csv": parse_csv,
}


def parse_document(file_path: str) -> tuple[str, list[str]]:
    """根据文件扩展名自动选择解析器，返回 (全文, 段落列表)"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in PARSERS:
        raise ValueError(f"不支持的格式: {ext}，支持: {', '.join(PARSERS.keys())}")
    return PARSERS[ext](file_path)


# ========== QA 解析器 ==========

QA_PATTERN = re.compile(
    r"^\s*(?:(?:问|Q|问题|Question)\s*[：:]\s*(.+?))\s*(?:\n|$)\s*"
    r"(?:(?:答|A|答案|Answer)\s*[：:]\s*(.+?))"
    r"(?=\n\s*(?:问|Q|问题|Question)\s*[：:]|\Z)",
    re.MULTILINE | re.DOTALL,
)


def detect_qa_format(text: str) -> bool:
    """检测文本是否为 QA 格式（包含 Q: / A: 或 问：/ 答： 配对）"""
    lines = text.strip().split("\n")
    q_count = 0
    a_count = 0
    for line in lines:
        stripped = line.strip()
        if re.match(r"^(问|Q|问题|Question)\s*[：:]\s*", stripped):
            q_count += 1
        elif re.match(r"^(答|A|答案|Answer)\s*[：:]\s*", stripped):
            a_count += 1
    return q_count >= 2 and a_count >= 2


def parse_qa(file_path: str) -> tuple[str, list[str]]:
    """
    解析 QA 格式文档
    返回 (全文, QA 块列表)，每个 QA 块是一个完整问答对
    """
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    pairs = QA_PATTERN.findall(text)
    chunks = []
    all_text_parts = []

    if pairs:
        for q, a in pairs:
            q = q.strip()
            a = a.strip()
            if q and a:
                block = f"问：{q}\n答：{a}"
                chunks.append(block)
                all_text_parts.append(f"Q: {q}\nA: {a}")
    else:
        lines = text.strip().split("\n")
        current_q = ""
        current_a = ""
        for line in lines:
            stripped = line.strip()
            qm = re.match(r"^(问|Q|问题|Question)\s*[：:]\s*(.+)", stripped)
            am = re.match(r"^(答|A|答案|Answer)\s*[：:]\s*(.+)", stripped)
            if qm:
                if current_q and current_a:
                    block = f"问：{current_q}\n答：{current_a}"
                    chunks.append(block)
                    all_text_parts.append(f"Q: {current_q}\nA: {current_a}")
                    current_a = ""
                current_q = qm.group(2).strip()
            elif am:
                current_a = am.group(2).strip()
                if current_q:
                    block = f"问：{current_q}\n答：{current_a}"
                    chunks.append(block)
                    all_text_parts.append(f"Q: {current_q}\nA: {current_a}")
                    current_q = ""
                    current_a = ""

        if current_q and current_a:
            block = f"问：{current_q}\n答：{current_a}"
            chunks.append(block)
            all_text_parts.append(f"Q: {current_q}\nA: {current_a}")

    if not chunks:
        raise ValueError("未能识别出 QA 问答对，请确认格式为 Q:/A: 或 问：/答：")

    return "\n\n".join(all_text_parts), chunks


def parse_with_type(file_path: str, force_type: str = "") -> tuple[str, list[str], str]:
    """
    根据文件类型或自动检测选择解析器
    返回 (全文, 块列表, 文档类型: "qa" 或 "doc")
    """
    if force_type == "qa":
        return *parse_qa(file_path), "qa"

    ext = os.path.splitext(file_path)[1].lower()
    raw_text, paragraphs = PARSERS[ext](file_path)

    if ext in (".md", ".txt") and detect_qa_format(raw_text):
        return *parse_qa(file_path), "qa"

    return raw_text, paragraphs, "doc"


# ========== 分块器 ==========

SENTENCE_END = re.compile(r"(?<=[。！？；\n.!?;])\s*")
TABLE_SEPARATOR = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+$")


def split_sentences(text: str) -> list[str]:
    """中英文句子分割"""
    parts = SENTENCE_END.split(text)
    return [p.strip() for p in parts if p.strip()]


def is_markdown_table(text: str) -> bool:
    """检测文本是否为 Markdown 表格格式"""
    lines = text.strip().split("\n")
    if len(lines) < 3:
        return False
    has_separator = False
    pipe_count = 0
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            pipe_count += 1
        if TABLE_SEPARATOR.match(stripped):
            has_separator = True
    return has_separator and pipe_count >= 3


def split_table_at_rows(table_text: str, max_size: int) -> list[str]:
    """在行边界切割过长的 Markdown 表格"""
    lines = table_text.strip().split("\n")
    if len(lines) < 3:
        return [table_text]

    header = lines[0]
    separator = lines[1]
    data_rows = lines[2:]

    chunks = []
    current_chunk = [header, separator]
    current_size = len(header) + len(separator) + 2

    for row in data_rows:
        row_size = len(row) + 1
        if current_size + row_size > max_size and len(current_chunk) > 2:
            chunks.append("\n".join(current_chunk))
            current_chunk = [header, separator]
            current_size = len(header) + len(separator) + 2
        current_chunk.append(row)
        current_size += row_size

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks


def chunk_document(paragraphs: list[str], chunk_size: int = CHUNK_SIZE,
                   overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    对段落列表进行分块
    - 尽量在段落边界分块
    - 超长段落按句子边界分割
    - 相邻块之间有 overlap 字符重叠
    """
    chunks = []
    buffer = ""

    for para in paragraphs:
        if len(buffer) + len(para) + 1 <= int(chunk_size * 1.1):
            if buffer:
                buffer += "\n\n" + para
            else:
                buffer = para
        else:
            if buffer:
                chunks.append(buffer)

            if len(para) > chunk_size:
                if is_markdown_table(para):
                    chunks.extend(split_table_at_rows(para, chunk_size))
                else:
                    sentences = split_sentences(para)
                    sub_buffer = ""
                    for sent in sentences:
                        if len(sub_buffer) + len(sent) + 1 <= chunk_size:
                            sub_buffer = (sub_buffer + "\n" + sent) if sub_buffer else sent
                        else:
                            if sub_buffer:
                                chunks.append(sub_buffer)
                            sub_buffer = sent
                    if sub_buffer:
                        chunks.append(sub_buffer)
                buffer = ""
            else:
                buffer = para

    if buffer:
        chunks.append(buffer)

    if overlap > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_end = chunks[i - 1]
            overlap_text = prev_end[-overlap:] if len(prev_end) > overlap else prev_end
            overlapped.append(overlap_text + "\n" + chunks[i])
        return overlapped

    return chunks
